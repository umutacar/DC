from __future__ import unicode_literals

import datetime

import sys
import re
import os
import json
import argparse
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

reload(sys)
sys.setdefaultencoding('utf8')

# local imports
import constants
BEGIN = "\\begin{"
END = "\\end{"
BEGIN_LEN = len("\\begin{")
END_LEN = len("\\end{")

TRUE_POS = 0
FALSE_NEG = 0
FALSE_POS = 0

def open_json_file(filename):
	with open(filename, "r") as read_file:
		print(read_file)
		data = json.load(read_file)
		return data

def file_to_dictionary_list(dict_filename):
	f = open(dict_filename, 'r')
	dict_str = f.read()
	lines = [d.split("\\label{") for d in dict_str.strip().split("\n")]
	lines = list(filter(lambda x: x, lines))
	# last element of dict_list is an empty list
	list_tuples = list(map(lambda x: (x[0].strip(), x[1].strip()[:-1]), lines))
	return list_tuples

def parse_atoms(tex):
	atoms = []
	tex = tex.decode().strip()
	# remove flex
	flex_instance = tex.find("\\begin{flex}")
	while (flex_instance != -1):
		label_index = tex.find("\\label{", flex_instance, -1)
		label_end_index = tex.find("}", label_index, -1)
		tex = tex[:flex_instance] + tex[label_end_index+1:]
		flex_end_index = tex.find("\\end{flex}")
		tex = tex[:flex_end_index] + tex[flex_end_index + len("\\end{flex}"):]
		flex_instance = tex.find("\\begin{flex}")
	# remove gram
	gram_instance = tex.find("\\begin{gram}")
	while (gram_instance != -1):
		label_index = tex.find("\\label{", gram_instance, -1)
		label_end_index = tex.find("}", label_index, -1)
		tex = tex[:gram_instance] + tex[label_end_index+1:]
		gram_end_index = tex.find("\\end{gram}")
		tex = tex[:gram_end_index] + tex[gram_end_index + len("\\end{gram}"):]
		gram_instance = tex.find("\\begin{gram}")
	while (tex):
		tex = tex.strip()
		begin_index = tex.find("\\begin{")
		# if (begin_index != -1):
		close_bracket = tex[begin_index:].index('}')
		arg = tex[begin_index + BEGIN_LEN: begin_index + close_bracket]
		end = tex.index("\\end{" + arg + "}")
		end_len = len("\\end{" + arg + "}")
		atom = tex[begin_index:end+end_len]
		atom = atom.split('\n')[1:-1] # remove begin and end
		if ("\\label{" in atom[0]):
			atom = atom[1:]
		atom = "\n".join(atom)
		atoms.append(atom)
		tex = tex[end+end_len:]
	return atoms

# Given: two lists of indices
def find_matches_per_keyphrase(keyterm, answers, predictions):
	a_i, p_i = 0, 0
	matches = []
	print(answers, predictions)
	while (a_i < len(answers) and p_i < len(predictions)):
		answer, pred = answers[a_i], predictions[p_i]
		# print(answer, pred)
		# pred : (index, length, phrase)
		if (answer[0] == pred[0]):
			# complete match, add as true positive
			if (answer[1] == pred[1]):
				matches.append( (keyterm, pred, 'G') )
				# TRUE_POS += 1
			# if the prediction term doesn't match, add as false positive
			else:
				matches.append( (keyterm, pred, 'R') )
				# FALSE_POS += 1
			a_i += 1
			p_i += 1
		# if there's an answer that was not detected, it's a false negative
		elif (answer[0] < pred[0]):
			matches.append((keyterm, answer, 'Y') )
			# FALSE_NEG += 1
			a_i += 1
		# if we detected something that's not an answer, it's false positive
		elif (answer[0] > pred[0]):
			matches.append((keyterm, pred, 'R') )
			# FALSE_POS += 1
			p_i += 1
	# false negatives
	while (a_i < len(answers)):
		matches.append((keyterm, answers[a_i], 'Y'))
		# FALSE_NEG += 1
		a_i += 1
	# false positives
	while (p_i < len(predictions)):
		matches.append( (keyterm, predictions[p_i], 'R') )
		# FALSE_POS += 1
		p_i += 1
	return matches

def find_matches_per_atom(answer_key, predictions):
	results = {}
	# find true positive and false negatives
	# pred_keyphrases = {}
	# for (prefix, suffix_dict) in predictions.items():
	# 	for (suffix, items) in suffix_dict.items():
	# 		pred_keyphrases[prefix + suffix] = [prefix, suffix]
	# keyphrases = list(set(answer_key.keys() + pred_keyphrases.keys()))
	keyphrases = list(set(answer_key.keys() + predictions.keys()))
	for keyphrase in keyphrases:
		true_indices, pred_indices = [], []
		if (keyphrase in answer_key):
			true_indices = answer_key[keyphrase]
		# if (keyphrase in pred_keyphrases.keys()):
		if (keyphrase in predictions):
			pred_indices = predictions[keyphrase]
			# prefix = pred_keyphrases[keyphrase][0]
			# suffix = pred_keyphrases[keyphrase][1]
			# pred_indices = predictions[prefix][suffix]

		if (true_indices == [] and pred_indices == []):
			continue
		# we predicted labels that don't exist
		elif (true_indices == []):
			results[keyphrase] = list(map(lambda x: (keyphrase, x, 'R'), pred_indices))
		# we didn't predict labels that are in the answer key
		elif (pred_indices == []):
			results[keyphrase] = list(map(lambda x: (keyphrase, x, 'Y'), true_indices))
		else:
			results[keyphrase] = find_matches_per_keyphrase(keyphrase, true_indices, pred_indices)
		print(results[keyphrase])
	return results

def generate_labelled_atom(doc, atom, matches):
	true_pos, false_pos, false_neg = 0, 0, 0
	atom = atom.decode()
	atom_index = 0
	p = doc.add_paragraph("")
	for i in range(len(matches)):
		(keyphrase, match_tuple, color) = matches[i]
		if (match_tuple == []):
			continue
		print(match_tuple)
		match_index, match_len = match_tuple[0], match_tuple[1]
		# decide color
		if color == 'R': 
			color_scheme = WD_COLOR_INDEX.RED 
			false_pos += 1
		elif color == 'Y':
			color_scheme = WD_COLOR_INDEX.YELLOW
			false_neg += 1
		elif color == 'G':
			color_scheme = WD_COLOR_INDEX.GREEN
			true_pos += 1

		match_end_index = match_index + match_len

		if (atom_index > match_index):
			atom_index = match_index
		unhighlighted_phrase = atom[atom_index:match_index]
		p.add_run(unhighlighted_phrase)
		highlighted_phrase = atom[match_index:match_end_index]
		# p.add_run(phrase).font.highlight_color = color_scheme
		p.add_run(highlighted_phrase).font.highlight_color = color_scheme

		atom_index = match_end_index
	p.add_run(atom[atom_index:])
	return (true_pos, false_pos, false_neg)


if __name__=='__main__':
	all_true_pos, all_false_pos, all_false_neg = 0, 0, 0
	parser = argparse.ArgumentParser()
	parser.add_argument('--answer_file', type=str, default='answer_key.json')
	parser.add_argument('--prediction_file', type=str)
	parser.add_argument('--tex_file', type=str)
	parser.add_argument('--output_file', type=str)
	args = parser.parse_args()
	
	f = open(args.tex_file, 'r')
	tex_str = f.read()
	atoms = parse_atoms(tex_str)
	answer_key = open_json_file(args.answer_file)
	experiment_results = open_json_file(args.prediction_file)
	doc = Document()
	doc.add_heading(args.prediction_file)
	for i in range(20):# len(atoms)):
		print("ATOM ", i)
		print(answer_key[i+10])
		print(len(experiment_results))
		results = find_matches_per_atom(answer_key[i+10], experiment_results[i])
		all_results = [] 
		for (keyterm, match_list) in results.items():
			all_results += match_list
		all_results.sort(key = lambda x: x[1]) # 0 -> keyphrase, 1 -> index, length, 2 -> color
		print(all_results)
		doc.add_heading('atom ' + str(i))
		(true_pos, false_pos, false_neg) = generate_labelled_atom(doc, atoms[i+10], all_results)
		all_true_pos += true_pos
		all_false_pos += false_pos
		all_false_neg += false_neg
	doc.save(args.output_file)
	print(all_true_pos, all_false_pos, all_false_neg)

